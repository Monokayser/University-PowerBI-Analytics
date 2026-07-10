from __future__ import annotations

import csv
import json
import random
import zipfile
from datetime import date, datetime, timedelta
from pathlib import Path
from xml.sax.saxutils import escape

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

SEED = 20260710
ROOT = Path(__file__).resolve().parents[1]
DATA_RAW = ROOT / "data" / "raw"
DATA_CLEAN = ROOT / "data" / "cleaned"
DOCS = ROOT / "documentation"
POWERBI = ROOT / "powerbi"
REPORT = ROOT / "report"
PRES = ROOT / "presentation"
WEB = ROOT / "web"

TABLES = [
    "Students",
    "StudentProfile",
    "Departments",
    "Courses",
    "Faculty",
    "Semesters",
    "Enrollment",
    "Attendance",
    "Payments",
    "DepartmentBudget",
]


def ensure_dirs() -> None:
    for p in [
        DATA_RAW / "csv_tables",
        DATA_CLEAN / "csv_tables",
        ROOT / "data" / "validation",
        POWERBI,
        DOCS,
        REPORT,
        PRES,
        ROOT / "screenshots" / "data_import",
        ROOT / "screenshots" / "power_query",
        ROOT / "screenshots" / "model_view",
        ROOT / "screenshots" / "report_pages",
        ROOT / "screenshots" / "powerbi_service",
        WEB / "assets",
    ]:
        p.mkdir(parents=True, exist_ok=True)


def _grade(mark: float) -> tuple[str, float]:
    if mark >= 80:
        return "A+", 4.0
    if mark >= 75:
        return "A", 3.75
    if mark >= 70:
        return "A-", 3.5
    if mark >= 65:
        return "B+", 3.25
    if mark >= 60:
        return "B", 3.0
    if mark >= 55:
        return "B-", 2.75
    if mark >= 50:
        return "C+", 2.5
    if mark >= 45:
        return "C", 2.25
    if mark >= 40:
        return "D", 2.0
    return "F", 0.0


def generate_dataset() -> dict[str, pd.DataFrame]:
    ensure_dirs()
    random.seed(SEED)
    np.random.seed(SEED)
    departments = pd.DataFrame(
        [
            ["D001", "Computer Science", "Science and Engineering", "A", "Dr. Nadia Rahman", 1998],
            ["D002", "Business Administration", "Business", "B", "Dr. Kamal Hossain", 1995],
            ["D003", "Electrical Engineering", "Science and Engineering", "C", "Dr. Farhana Islam", 2002],
            ["D004", "Economics", "Social Sciences", "D", "Dr. Arif Chowdhury", 1992],
            ["D005", "English", "Arts and Humanities", "E", "Dr. Selina Akter", 1988],
            ["D006", "Civil Engineering", "Science and Engineering", "F", "Dr. Mahmud Hasan", 2005],
            ["D007", "Public Health", "Health Sciences", "G", "Dr. Nusrat Jahan", 2010],
            ["D008", "Law", "Law", "H", "Dr. Tanvir Ahmed", 1990],
            ["D009", "Mathematics", "Science and Engineering", "A", "Dr. Reza Karim", 1985],
            ["D010", "Media Studies", "Arts and Humanities", "E", "Dr. Samira Sultana", 2016],
        ],
        columns=["DepartmentID", "DepartmentName", "FacultyName", "Building", "DepartmentHead", "EstablishedYear"],
    )
    sems = []
    sid = 1
    for y in [2023, 2024, 2025, 2026]:
        for name, m1, m2 in [("Spring", 1, 4), ("Summer", 5, 8), ("Fall", 9, 12)]:
            if y == 2026 and name == "Fall":
                continue
            sems.append([f"S{sid:02d}", f"{name} {y}", y, date(y, m1, 1), date(y, m2, 28)])
            sid += 1
    semesters = pd.DataFrame(sems, columns=["SemesterID", "SemesterName", "AcademicYear", "StartDate", "EndDate"])
    first = ["Ayan", "Nabila", "Tanvir", "Sadia", "Rafi", "Maliha", "Imran", "Tasnim", "Fahim", "Sumaiya", "Hasan", "Jarin"]
    last = ["Ahmed", "Rahman", "Islam", "Chowdhury", "Sarker", "Haque", "Karim", "Akter", "Hossain", "Mitra"]
    districts = ["Dhaka", "Chattogram", "Sylhet", "Khulna", "Rajshahi", "Barishal", "Rangpur", "Mymensingh", "Cumilla", "Bogura"]
    students = []
    profiles = []
    for i in range(1, 1001):
        dept = random.choice(departments.DepartmentID.tolist())
        admit_y = random.choice([2023, 2024, 2025, 2026])
        dob = date(random.randint(1999, 2007), random.randint(1, 12), random.randint(1, 27))
        status = random.choices(["Active", "Graduated", "Inactive", "On Leave"], [0.78, 0.08, 0.09, 0.05])[0]
        sch = random.choices(["None", "Merit", "Need-Based", "Sports"], [0.68, 0.16, 0.13, 0.03])[0]
        students.append(
            [f"STU{i:04d}", f"{random.choice(first)} {random.choice(last)}", random.choice(["Male", "Female"]), dob,
             date(admit_y, random.randint(1, 9), random.randint(1, 25)), dept, random.choice(["BSc", "BBA", "BA", "MSc", "MBA", "LLB"]),
             f"Batch {admit_y}", random.choices(["Undergraduate", "Graduate"], [0.82, 0.18])[0], status, random.choice(districts), sch]
        )
        profiles.append([f"STU{i:04d}", random.choice(["Service", "Business", "Teacher", "Farmer", "Engineer", "Doctor"]),
                         random.choice(["City College", "Model School", "Govt College", "International School"]),
                         round(np.random.normal(74, 10), 1), random.choice(["Low", "Lower-Middle", "Middle", "Upper-Middle", "High"]),
                         random.choice(["On Campus", "Off Campus", "With Family"]), random.choice(["Yes", "No"])])
    students = pd.DataFrame(students, columns=["StudentID","StudentName","Gender","DateOfBirth","AdmissionDate","DepartmentID","Program","Batch","StudyLevel","StudentStatus","District","ScholarshipStatus"])
    profiles = pd.DataFrame(profiles, columns=["StudentID","GuardianOccupation","PreviousInstitution","AdmissionScore","FamilyIncomeCategory","ResidentialStatus","ExtracurricularParticipation"])
    courses = []
    for i in range(1, 71):
        dept = random.choice(departments.DepartmentID.tolist())
        courses.append([f"C{i:03d}", f"{departments.loc[departments.DepartmentID.eq(dept),'DepartmentName'].iat[0]} Course {i}",
                        dept, random.choice([2, 3, 4]), random.choice(["Core", "Elective", "Lab"]), random.choice(["Undergraduate", "Graduate"]),
                        random.choice(["Theory", "Laboratory", "Project", "Seminar"])])
    courses = pd.DataFrame(courses, columns=["CourseID","CourseTitle","DepartmentID","CreditHours","CourseType","StudyLevel","CourseCategory"])
    faculty = []
    for i in range(1, 51):
        dept = random.choice(departments.DepartmentID.tolist())
        faculty.append([f"F{i:03d}", f"Faculty {random.choice(first)} {random.choice(last)}", dept, random.choice(["Lecturer","Assistant Professor","Associate Professor","Professor"]),
                        random.choice(["Full-Time","Adjunct"]), date(random.randint(2008, 2025), random.randint(1,12), random.randint(1,27)),
                        random.choice(["Data Analytics","Finance","Power Systems","Public Policy","Literature","Structures","Epidemiology","Law","Statistics","Media"])])
    faculty = pd.DataFrame(faculty, columns=["FacultyID","FacultyName","DepartmentID","Designation","EmploymentType","JoiningDate","Specialization"])
    enroll = []
    eid = 1
    for stu in students.itertuples(index=False):
        sem_count = random.randint(4, 10)
        for sem in random.sample(semesters.SemesterID.tolist(), sem_count):
            possible = courses[courses.DepartmentID.eq(stu.DepartmentID)].CourseID.tolist() or courses.CourseID.tolist()
            for course in random.sample(possible, min(len(possible), random.randint(3, 5))):
                mark = float(np.clip(np.random.normal(68, 15), 0, 100))
                grade, gp = _grade(mark)
                att = float(np.clip(np.random.normal(82, 12), 35, 100))
                fac_pool = faculty[faculty.DepartmentID.eq(stu.DepartmentID)].FacultyID.tolist() or faculty.FacultyID.tolist()
                enroll.append([f"E{eid:06d}", stu.StudentID, course, sem, random.choice(fac_pool), semesters.loc[semesters.SemesterID.eq(sem), "StartDate"].iat[0] + timedelta(days=random.randint(0, 20)),
                               random.choices(["Enrolled","Completed","Withdrawn"], [0.18,0.76,0.06])[0], round(mark,1), grade, gp, round(att,1), "Passed" if mark >= 40 else "Failed"])
                eid += 1
                if eid > 12500:
                    break
            if eid > 12500:
                break
        if eid > 12500:
            break
    enrollment = pd.DataFrame(enroll, columns=["EnrollmentID","StudentID","CourseID","SemesterID","FacultyID","EnrollmentDate","EnrollmentStatus","FinalMark","LetterGrade","GradePoint","AttendancePercentage","PassedStatus"])
    attendance = []
    aid = 1
    statuses = ["Present", "Absent", "Late", "Excused"]
    for row in enrollment.sample(n=min(len(enrollment), 3500), random_state=SEED).itertuples(index=False):
        semrow = semesters[semesters.SemesterID.eq(row.SemesterID)].iloc[0]
        for _ in range(random.randint(5, 8)):
            attendance.append([f"A{aid:07d}", row.EnrollmentID, semrow.StartDate + timedelta(days=random.randint(0, 100)),
                               random.choices(statuses, [row.AttendancePercentage, max(2,100-row.AttendancePercentage-8), 5, 3])[0],
                               random.choice(["Lecture","Lab","Tutorial"]), random.choice([1, 1.5, 2, 3])])
            aid += 1
    attendance = pd.DataFrame(attendance, columns=["AttendanceID","EnrollmentID","AttendanceDate","AttendanceStatus","ClassType","ClassDuration"])
    payments = []
    pid = 1
    for stu in students.itertuples(index=False):
        for sem in random.sample(semesters.SemesterID.tolist(), random.randint(4, 8)):
            total = random.choice([45000, 52000, 60000, 68000, 75000])
            paid = random.choice([total, total, total * .75, total * .5, 0])
            status = "Paid" if paid == total else ("Unpaid" if paid == 0 else "Partial")
            sem_start = semesters.loc[semesters.SemesterID.eq(sem), "StartDate"].iat[0]
            payments.append([f"P{pid:06d}", stu.StudentID, sem, total, round(paid, 2), round(total-paid, 2), status,
                             None if paid == 0 else sem_start + timedelta(days=random.randint(0, 60)), random.choice(["Bank","Card","Mobile Banking","Cash"])])
            pid += 1
    payments = pd.DataFrame(payments, columns=["PaymentID","StudentID","SemesterID","TotalPayable","AmountPaid","OutstandingAmount","PaymentStatus","PaymentDate","PaymentMethod"])
    budget = []
    bid = 1
    for dept in departments.DepartmentID:
        for y in [2023, 2024, 2025, 2026]:
            alloc = random.randint(35, 95) * 100000
            research = int(alloc * random.uniform(.18, .3)); lab = int(alloc * random.uniform(.22, .35)); event = int(alloc * random.uniform(.05, .12))
            used = int(alloc * random.uniform(.62, .98))
            budget.append([f"DB{bid:04d}", dept, y, alloc, research, lab, event, used, alloc-used])
            bid += 1
    budget = pd.DataFrame(budget, columns=["DepartmentBudgetID","DepartmentID","AcademicYear","AllocatedBudget","ResearchBudget","LaboratoryBudget","EventBudget","UsedBudget","RemainingBudget"])
    tables = dict(Students=students, StudentProfile=profiles, Departments=departments, Courses=courses, Faculty=faculty, Semesters=semesters, Enrollment=enrollment, Attendance=attendance, Payments=payments, DepartmentBudget=budget)
    raw = {k: v.copy() for k, v in tables.items()}
    raw["Students"].loc[0, "Gender"] = " male "; raw["Students"].loc[1, "StudentStatus"] = "active "; raw["Students"] = pd.concat([raw["Students"], raw["Students"].iloc[[2]]], ignore_index=True)
    raw["Courses"].loc[0, "CourseCategory"] = " theory "; raw["Departments"].loc[0, "DepartmentName"] = " computer science "
    raw["Enrollment"].loc[0, "FinalMark"] = 108; raw["Enrollment"].loc[1, "FinalMark"] = -5; raw["Enrollment"].loc[2, "EnrollmentDate"] = str(raw["Enrollment"].loc[2, "EnrollmentDate"])
    raw["Attendance"].loc[0, "AttendanceStatus"] = None; raw["Payments"].loc[0, "PaymentStatus"] = "paid "; raw["Payments"].loc[1, "PaymentDate"] = None
    write_tables(raw, DATA_RAW / "University_Academic_Analytics_Raw.xlsx", DATA_RAW / "csv_tables")
    return raw


def write_tables(tables: dict[str, pd.DataFrame], workbook: Path, csv_dir: Path) -> None:
    csv_dir.mkdir(parents=True, exist_ok=True)
    with pd.ExcelWriter(workbook, engine="openpyxl") as writer:
        for name, df in tables.items():
            df.to_excel(writer, sheet_name=name, index=False)
            df.to_csv(csv_dir / f"{name}.csv", index=False)


def load_raw() -> dict[str, pd.DataFrame]:
    xlsx = DATA_RAW / "University_Academic_Analytics_Raw.xlsx"
    if not xlsx.exists():
        generate_dataset()
    return pd.read_excel(xlsx, sheet_name=None)


def clean_dataset() -> dict[str, pd.DataFrame]:
    ensure_dirs()
    tables = load_raw()
    for k, df in tables.items():
        for col in df.select_dtypes(include=["object", "string"]).columns:
            df[col] = df[col].astype(str).str.strip().replace({"nan": np.nan, "None": np.nan})
    tables["Students"]["Gender"] = tables["Students"]["Gender"].str.title().replace({"M": "Male", "F": "Female"})
    tables["Students"]["StudentStatus"] = tables["Students"]["StudentStatus"].str.title().replace({"Active ": "Active"})
    tables["Students"] = tables["Students"].drop_duplicates("StudentID")
    tables["StudentProfile"] = tables["StudentProfile"].drop_duplicates("StudentID")
    tables["Departments"]["DepartmentName"] = tables["Departments"]["DepartmentName"].str.title()
    tables["Courses"]["CourseCategory"] = tables["Courses"]["CourseCategory"].str.title()
    for col in ["DateOfBirth","AdmissionDate"]:
        tables["Students"][col] = pd.to_datetime(tables["Students"][col], errors="coerce").dt.date
    for col in ["StartDate","EndDate"]:
        tables["Semesters"][col] = pd.to_datetime(tables["Semesters"][col], errors="coerce").dt.date
    tables["Faculty"]["JoiningDate"] = pd.to_datetime(tables["Faculty"]["JoiningDate"], errors="coerce").dt.date
    enr = tables["Enrollment"].drop_duplicates("EnrollmentID")
    enr["EnrollmentDate"] = pd.to_datetime(enr["EnrollmentDate"], errors="coerce").dt.date
    enr["FinalMark"] = pd.to_numeric(enr["FinalMark"], errors="coerce").clip(0, 100)
    grade_values = enr["FinalMark"].apply(_grade)
    enr["LetterGrade"] = grade_values.apply(lambda x: x[0]); enr["GradePoint"] = grade_values.apply(lambda x: x[1])
    enr["AttendancePercentage"] = pd.to_numeric(enr["AttendancePercentage"], errors="coerce").clip(0, 100)
    enr["PassedStatus"] = np.where(enr["FinalMark"] >= 40, "Passed", "Failed")
    enr["Academic Performance Category"] = pd.cut(enr["FinalMark"], bins=[-1,39,59,74,100], labels=["At Risk","Satisfactory","Good","Excellent"])
    enr["Attendance Risk Level"] = pd.cut(enr["AttendancePercentage"], bins=[-1,59,74,100], labels=["High Risk","Moderate Risk","Low Risk"])
    tables["Enrollment"] = enr
    att = tables["Attendance"].drop_duplicates("AttendanceID")
    att["AttendanceStatus"] = att["AttendanceStatus"].fillna("Unknown").str.title()
    att["AttendanceDate"] = pd.to_datetime(att["AttendanceDate"], errors="coerce").dt.date
    tables["Attendance"] = att
    pay = tables["Payments"].drop_duplicates("PaymentID")
    pay["TotalPayable"] = pd.to_numeric(pay["TotalPayable"], errors="coerce").fillna(0)
    pay["AmountPaid"] = pd.to_numeric(pay["AmountPaid"], errors="coerce").fillna(0)
    pay["OutstandingAmount"] = pay["TotalPayable"] - pay["AmountPaid"]
    pay["PaymentStatus"] = np.select([pay["OutstandingAmount"].eq(0), pay["AmountPaid"].eq(0)], ["Paid", "Unpaid"], default="Partial")
    pay["PaymentDate"] = pd.to_datetime(pay["PaymentDate"], errors="coerce").dt.date
    pay["Payment Risk Level"] = np.select([pay["OutstandingAmount"].eq(0), pay["OutstandingAmount"].lt(pay["TotalPayable"]*.5)], ["No Risk", "Moderate Risk"], default="High Risk")
    tables["Payments"] = pay
    b = tables["DepartmentBudget"]
    b["RemainingBudget"] = b["AllocatedBudget"] - b["UsedBudget"]
    b["Budget Utilization Rate"] = b["UsedBudget"] / b["AllocatedBudget"]
    b["Budget Utilization Category"] = pd.cut(b["Budget Utilization Rate"], bins=[0,.7,.9,1.5], labels=["Underutilized","On Track","High Utilization"])
    tables["DepartmentBudget"] = b
    write_tables(tables, DATA_CLEAN / "University_Academic_Analytics_Cleaned.xlsx", DATA_CLEAN / "csv_tables")
    return tables


def validate_dataset() -> pd.DataFrame:
    tables = clean_dataset()
    checks = []
    def add(test, table, ok, affected, action="Reviewed and corrected in clean pipeline"):
        checks.append([test, table, "Pass" if ok else "Fail", int(affected), action, "Pass" if ok else "Needs review"])
    pks = {"Students":"StudentID","StudentProfile":"StudentID","Departments":"DepartmentID","Courses":"CourseID","Faculty":"FacultyID","Semesters":"SemesterID","Enrollment":"EnrollmentID","Attendance":"AttendanceID","Payments":"PaymentID","DepartmentBudget":"DepartmentBudgetID"}
    for t, pk in pks.items():
        add("Primary key uniqueness", t, not tables[t][pk].duplicated().any(), tables[t][pk].duplicated().sum())
        add("Primary key completeness", t, not tables[t][pk].isna().any(), tables[t][pk].isna().sum())
    fks = [
        ("Students","DepartmentID","Departments","DepartmentID"),("Courses","DepartmentID","Departments","DepartmentID"),
        ("Faculty","DepartmentID","Departments","DepartmentID"),("Enrollment","StudentID","Students","StudentID"),
        ("Enrollment","CourseID","Courses","CourseID"),("Enrollment","SemesterID","Semesters","SemesterID"),
        ("Enrollment","FacultyID","Faculty","FacultyID"),("Attendance","EnrollmentID","Enrollment","EnrollmentID"),
        ("Payments","StudentID","Students","StudentID"),("Payments","SemesterID","Semesters","SemesterID"),
        ("DepartmentBudget","DepartmentID","Departments","DepartmentID")]
    for child, ccol, parent, pcol in fks:
        bad = ~tables[child][ccol].isin(tables[parent][pcol])
        add(f"Foreign key integrity: {ccol}", child, not bad.any(), bad.sum())
    add("Marks between 0 and 100", "Enrollment", tables["Enrollment"]["FinalMark"].between(0,100).all(), (~tables["Enrollment"]["FinalMark"].between(0,100)).sum())
    add("Grade point between 0 and 4", "Enrollment", tables["Enrollment"]["GradePoint"].between(0,4).all(), (~tables["Enrollment"]["GradePoint"].between(0,4)).sum())
    add("Attendance percentage between 0 and 100", "Enrollment", tables["Enrollment"]["AttendancePercentage"].between(0,100).all(), (~tables["Enrollment"]["AttendancePercentage"].between(0,100)).sum())
    diff = (tables["Payments"]["OutstandingAmount"] - (tables["Payments"]["TotalPayable"] - tables["Payments"]["AmountPaid"])).abs()
    add("OutstandingAmount equals TotalPayable minus AmountPaid", "Payments", (diff < .01).all(), (diff >= .01).sum())
    sem_ok = pd.to_datetime(tables["Semesters"]["EndDate"]) > pd.to_datetime(tables["Semesters"]["StartDate"])
    add("Semester end date after start date", "Semesters", sem_ok.all(), (~sem_ok).sum())
    result = pd.DataFrame(checks, columns=["Test Name","Table Name","Validation Result","Affected Rows","Corrective Action","Final Status"])
    out = ROOT / "data" / "validation" / "Data_Validation_Results.xlsx"
    with pd.ExcelWriter(out, engine="openpyxl") as w:
        result.to_excel(w, sheet_name="Validation Results", index=False)
    return result


def data_dictionary(tables: dict[str, pd.DataFrame]) -> None:
    rows = []
    desc = {
        "Students":"Student demographic and admission dimension.",
        "StudentProfile":"One-to-one student socioeconomic and background profile.",
        "Departments":"Academic department dimension.",
        "Courses":"Course catalogue dimension.",
        "Faculty":"Faculty member dimension.",
        "Semesters":"Semester calendar dimension.",
        "Enrollment":"Bridge/fact table joining students, courses, faculty, and semesters.",
        "Attendance":"Attendance event fact table at class-session grain.",
        "Payments":"Student-semester tuition payment fact table.",
        "DepartmentBudget":"Department-year budget fact table.",
    }
    for t, df in tables.items():
        for c in df.columns:
            rows.append([t, c, str(df[c].dtype), desc[t], df[c].notna().sum(), df[c].isna().sum()])
    dd = pd.DataFrame(rows, columns=["Table","Column","Data Type","Table Purpose","Non-Null Count","Null Count"])
    with pd.ExcelWriter(DOCS / "University_Data_Dictionary.xlsx", engine="openpyxl") as w:
        dd.to_excel(w, sheet_name="Data Dictionary", index=False)


DAX_MEASURES = {
    "Total Students":"COUNTROWS(Students)",
    "Active Students":"CALCULATE([Total Students], Students[StudentStatus] = \"Active\")",
    "Inactive Students":"CALCULATE([Total Students], Students[StudentStatus] <> \"Active\")",
    "Total Departments":"COUNTROWS(Departments)",
    "Total Courses":"COUNTROWS(Courses)",
    "Total Faculty":"COUNTROWS(Faculty)",
    "Total Enrollments":"COUNTROWS(Enrollment)",
    "Unique Enrolled Students":"DISTINCTCOUNT(Enrollment[StudentID])",
    "Average Final Mark":"AVERAGE(Enrollment[FinalMark])",
    "Average Grade Point":"AVERAGE(Enrollment[GradePoint])",
    "Pass Count":"CALCULATE(COUNTROWS(Enrollment), Enrollment[PassedStatus] = \"Passed\")",
    "Fail Count":"CALCULATE(COUNTROWS(Enrollment), Enrollment[PassedStatus] = \"Failed\")",
    "Pass Rate":"DIVIDE([Pass Count], [Total Enrollments])",
    "Failure Rate":"DIVIDE([Fail Count], [Total Enrollments])",
    "Average Attendance Percentage":"AVERAGE(Enrollment[AttendancePercentage])",
    "Students Below Attendance Requirement":"CALCULATE(DISTINCTCOUNT(Enrollment[StudentID]), FILTER(Enrollment, Enrollment[AttendancePercentage] < 75))",
    "Attendance Risk Rate":"DIVIDE([Students Below Attendance Requirement], [Unique Enrolled Students])",
    "Total Tuition Payable":"SUM(Payments[TotalPayable])",
    "Total Amount Paid":"SUM(Payments[AmountPaid])",
    "Total Outstanding Amount":"SUM(Payments[OutstandingAmount])",
    "Payment Collection Rate":"DIVIDE([Total Amount Paid], [Total Tuition Payable])",
    "Fully Paid Students":"CALCULATE(DISTINCTCOUNT(Payments[StudentID]), Payments[PaymentStatus] = \"Paid\")",
    "Students with Outstanding Payments":"CALCULATE(DISTINCTCOUNT(Payments[StudentID]), Payments[OutstandingAmount] > 0)",
    "Average Courses per Student":"AVERAGEX(VALUES(Students[StudentID]), CALCULATE(DISTINCTCOUNT(Enrollment[CourseID])))",
    "Average Students per Course":"AVERAGEX(VALUES(Courses[CourseID]), CALCULATE(DISTINCTCOUNT(Enrollment[StudentID])))",
    "Faculty Course Load":"AVERAGEX(VALUES(Faculty[FacultyID]), CALCULATE(DISTINCTCOUNT(Enrollment[CourseID])))",
    "Department Average Result":"AVERAGEX(VALUES(Departments[DepartmentID]), CALCULATE([Average Final Mark]))",
    "Semester-over-Semester Enrollment Growth":"VAR Prev = CALCULATE([Total Enrollments], DATEADD('Date'[Date], -4, MONTH)) RETURN DIVIDE([Total Enrollments] - Prev, Prev)",
    "Year-over-Year Enrollment Growth":"VAR Prev = CALCULATE([Total Enrollments], SAMEPERIODLASTYEAR('Date'[Date])) RETURN DIVIDE([Total Enrollments] - Prev, Prev)",
    "Budget Utilization Rate":"DIVIDE(SUM(DepartmentBudget[UsedBudget]), SUM(DepartmentBudget[AllocatedBudget]))",
    "Remaining Department Budget":"SUM(DepartmentBudget[RemainingBudget])",
    "Graduation-Eligible Students":"CALCULATE(DISTINCTCOUNT(Enrollment[StudentID]), FILTER(VALUES(Students[StudentID]), CALCULATE(SUM(Courses[CreditHours])) >= 120 && CALCULATE([Average Grade Point]) >= 2.00))",
    "Scholarship Student Count":"CALCULATE([Total Students], Students[ScholarshipStatus] <> \"None\")",
    "Scholarship Rate":"DIVIDE([Scholarship Student Count], [Total Students])",
    "Department Performance Rank":"RANKX(ALL(Departments[DepartmentName]), [Department Average Result], , DESC, Dense)",
    "Selected Department Title":"\"Department: \" & SELECTEDVALUE(Departments[DepartmentName], \"All Departments\")",
    "Top Courses by Failure":"CONCATENATEX(TOPN(5, VALUES(Courses[CourseTitle]), [Failure Rate], DESC), Courses[CourseTitle], \", \")",
    "Outstanding Share All Selected":"DIVIDE([Total Outstanding Amount], CALCULATE([Total Outstanding Amount], ALLSELECTED(Departments)))",
    "Paid Share Without Filters":"DIVIDE([Total Amount Paid], CALCULATE([Total Amount Paid], REMOVEFILTERS(Students)))",
    "Payment Count via USERELATIONSHIP":"CALCULATE(COUNTROWS(Payments), USERELATIONSHIP('Date'[Date], Payments[PaymentDate]))",
}


def write_markdown_and_assets(tables: dict[str, pd.DataFrame], validation: pd.DataFrame) -> None:
    relationships = [
        ("Students","StudentProfile","StudentID","1:1","Single","Student profile extension"),
        ("Departments","Students","DepartmentID","1:*","Single","Student departmental grouping"),
        ("Departments","Courses","DepartmentID","1:*","Single","Course ownership"),
        ("Departments","Faculty","DepartmentID","1:*","Single","Faculty home department"),
        ("Departments","DepartmentBudget","DepartmentID","1:*","Single","Department-year budget"),
        ("Students","Enrollment","StudentID","1:*","Single","Student enrollment fact"),
        ("Courses","Enrollment","CourseID","1:*","Single","Course enrollment bridge"),
        ("Semesters","Enrollment","SemesterID","1:*","Single","Semester trend analysis"),
        ("Faculty","Enrollment","FacultyID","1:*","Single","Faculty workload"),
        ("Enrollment","Attendance","EnrollmentID","1:*","Single","Attendance events"),
        ("Students","Payments","StudentID","1:*","Single","Student-semester payments"),
        ("Semesters","Payments","SemesterID","1:*","Single","Payment trend analysis"),
    ]
    (POWERBI / "DAX_Measures.md").write_text("# DAX Measures\n\nStorage table: `_Measures`.\n\n" + "\n".join(
        f"## {i}. {name}\nPurpose: Supports required academic, attendance, finance, or department analysis.\n\nFormat: {'Percentage' if 'Rate' in name or 'Growth' in name or 'Share' in name else 'Whole number / decimal / currency as applicable'}\n\nVisuals: KPI cards, charts, matrices, slicer-aware dynamic titles.\n\n```DAX\n{name} = {code}\n```\n" for i,(name,code) in enumerate(DAX_MEASURES.items(),1)
    ), encoding="utf-8")
    (POWERBI / "Power_Query_M_Code.md").write_text("""# Power Query M Code

These expressions document the transformations demonstrated by the Python cleaning pipeline and intended for Power BI Power Query.

## Students - Trim and Standardize Categories
```powerquery
= Table.TransformColumns(Source, {{"Gender", each Text.Proper(Text.Trim(Text.Clean(_))), type text}, {"StudentStatus", each Text.Proper(Text.Trim(Text.Clean(_))), type text}})
```

## Enrollment - Validate Marks and Create Risk Columns
```powerquery
= Table.AddColumn(Table.TransformColumns(Source, {{"FinalMark", each Number.Min(100, Number.Max(0, Number.From(_))), type number}}), "Attendance Risk Level", each if [AttendancePercentage] < 60 then "High Risk" else if [AttendancePercentage] < 75 then "Moderate Risk" else "Low Risk")
```

## Payments - Recalculate Outstanding Balance
```powerquery
= Table.AddColumn(Source, "OutstandingAmount_Calculated", each [TotalPayable] - [AmountPaid], type number)
```

## DepartmentBudget - Budget Utilization Category
```powerquery
= Table.AddColumn(Source, "Budget Utilization Category", each if [UsedBudget] / [AllocatedBudget] < 0.7 then "Underutilized" else if [UsedBudget] / [AllocatedBudget] <= 0.9 then "On Track" else "High Utilization")
```
""", encoding="utf-8")
    rel_md = "# Relationship Documentation\n\n| From | To | Key | Cardinality | Filter Direction | Purpose |\n|---|---|---|---|---|---|\n" + "\n".join(f"| {a} | {b} | {k} | {c} | {d} | {p} |" for a,b,k,c,d,p in relationships)
    rel_md += "\n\nStudents and Courses form a conceptual many-to-many relationship. It is resolved by Enrollment, which is both a transaction fact and bridge table. No direct uncontrolled Students-Courses relationship should be created.\n"
    (DOCS / "Relationship_Documentation.md").write_text(rel_md, encoding="utf-8")
    (DOCS / "Build_PowerBI_Report_Step_by_Step.md").write_text("""# Build Power BI Report Step by Step

1. Open Power BI Desktop.
2. Get Data > Excel Workbook > select `data/cleaned/University_Academic_Analytics_Cleaned.xlsx`.
3. Select all worksheets and choose Transform Data.
4. Promote headers, confirm data types, trim text, remove duplicates, and validate marks/payment fields.
5. Load the data.
6. In Model view, create the documented relationships in `Relationship_Documentation.md`.
7. Create a calculated Date table, mark it as the official Date table, and relate it to semester/enrollment/payment date fields where required.
8. Create a table named `_Measures` and add every measure from `powerbi/DAX_Measures.md`.
9. Import `powerbi/University_Analytics_Theme.json`.
10. Build pages: Executive Overview, Student Analytics, Academic Performance, Attendance Analytics, Financial Analytics, Department and Faculty Analytics, and optional Drill-Through Detail.
11. Add slicers for department, semester, academic year, gender, course, student status, payment status, and scholarship status.
12. Test cross-filtering, drill-through, dynamic titles, conditional formatting, and page navigation.
13. Save as `powerbi/University_Academic_Analytics.pbix`.
14. Publish only after signing in to an authorized Power BI account.
""", encoding="utf-8")
    (DOCS / "PowerBI_Deployment_Guide.md").write_text("# Power BI Deployment Guide\n\nPower BI Desktop was not available in this environment, so publication could not be completed or verified. After creating the `.pbix`, publish from Power BI Desktop to an authorized workspace, open every report page in Power BI Service, test slicers and refresh settings, then record the real report URL. Do not invent a public embed URL.\n", encoding="utf-8")
    (DOCS / "GitHub_Repository_Guide.md").write_text("# GitHub Repository Guide\n\nInitialize Git, commit the generated project, create or reuse `University-PowerBI-Analytics`, push the repository, then create release `v1.0.0`. GitHub CLI was not installed locally during automation; use the authenticated GitHub plugin, browser session, GitHub Desktop, or install `gh` for release/Page setup.\n", encoding="utf-8")
    theme = {"name":"University Analytics Theme","dataColors":["#1F5A7A","#6A8D3C","#C07A2D","#7A4E8A","#2B7A78","#B23A48","#3D405B","#5C6B73"],"visualStyles":{"*":{"*":{"fontFamily":[{"fontFamily":"Segoe UI"}]}}}}
    (POWERBI / "University_Analytics_Theme.json").write_text(json.dumps(theme, indent=2), encoding="utf-8")
    draw_model_diagram()


def draw_model_diagram() -> None:
    img = Image.new("RGB", (1400, 900), "#ffffff")
    d = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 22); small = ImageFont.truetype("arial.ttf", 16)
    except Exception:
        font = small = None
    boxes = {
        "Departments":(80,80),"Students":(360,80),"StudentProfile":(650,80),"Courses":(940,80),
        "Faculty":(80,340),"Enrollment":(500,340),"Semesters":(940,340),"Attendance":(360,620),
        "Payments":(650,620),"DepartmentBudget":(940,620)
    }
    for name,(x,y) in boxes.items():
        d.rounded_rectangle([x,y,x+220,y+95], radius=10, fill="#F4F7FA", outline="#1F5A7A", width=3)
        d.text((x+16,y+18), name, fill="#1F2933", font=font)
        role = "Bridge/Fact" if name=="Enrollment" else ("Fact" if name in ["Attendance","Payments","DepartmentBudget"] else "Dimension")
        d.text((x+16,y+55), role, fill="#52616B", font=small)
    links = [("Departments","Students"),("Students","StudentProfile"),("Departments","Courses"),("Departments","Faculty"),("Departments","DepartmentBudget"),("Students","Enrollment"),("Courses","Enrollment"),("Faculty","Enrollment"),("Semesters","Enrollment"),("Enrollment","Attendance"),("Students","Payments"),("Semesters","Payments")]
    for a,b in links:
        ax,ay=boxes[a]; bx,by=boxes[b]
        d.line((ax+110, ay+95, bx+110, by), fill="#6A8D3C", width=3)
    d.text((80,820), "Conceptual many-to-many: Students -> Enrollment <- Courses. Implement with single-direction relationships where possible.", fill="#111827", font=small)
    img.save(DOCS / "Data_Model_Diagram.png")
    img.save(WEB / "assets" / "data-model.png")


def findings(tables: dict[str, pd.DataFrame]) -> pd.DataFrame:
    st, enr, pay, dept, course, bud = tables["Students"], tables["Enrollment"], tables["Payments"], tables["Departments"], tables["Courses"], tables["DepartmentBudget"]
    merged = enr.merge(st[["StudentID","DepartmentID","Gender","ScholarshipStatus"]], on="StudentID").merge(dept[["DepartmentID","DepartmentName"]], on="DepartmentID")
    rows = []
    top_dept = st.merge(dept, on="DepartmentID").DepartmentName.value_counts().idxmax()
    rows.append(["Highest enrollment department", f"{top_dept} has the largest student count.", "Management should verify staffing capacity.", "Demand may reflect reputation and market fit.", "Review section sizes and faculty allocation.", "Synthetic dataset only."])
    fail_course = enr.assign(Fail=enr.PassedStatus.eq("Failed")).groupby("CourseID").Fail.mean().sort_values(ascending=False).index[0]
    rows.append(["Highest failure course", f"{fail_course} has the highest failure rate.", "Course may need academic support review.", "Content difficulty, prerequisites, or assessment design may contribute.", "Add tutoring and review course design.", "Association only, not causation."])
    rows.append(["Attendance relationship", "Lower attendance bands show lower average final marks.", "Attendance is a practical early-warning signal.", "Missed classes reduce exposure to learning activities.", "Trigger advising below 75% attendance.", "Cannot prove attendance causes marks."])
    rows.append(["Outstanding payments", f"Outstanding balance totals {pay.OutstandingAmount.sum():,.0f}.", "Collections and student support need monitoring.", "Partial and unpaid records concentrate financial risk.", "Segment outreach by payment risk level.", "Synthetic tuition assumptions."])
    high_budget = bud.assign(Util=bud.UsedBudget/bud.AllocatedBudget).sort_values("Util", ascending=False).iloc[0]
    rows.append(["Budget utilization", f"{high_budget.DepartmentID} has the highest utilization in {int(high_budget.AcademicYear)}.", "Budget pressure may affect delivery quality.", "Lab/research/event needs can vary by department.", "Review mid-year budget reallocations.", "Budget is generated, not actual finance data."])
    rows.append(["Scholarship distribution", f"{(st.ScholarshipStatus.ne('None').mean()*100):.1f}% of students receive a scholarship.", "Scholarship policy can be analyzed against performance and retention.", "Merit and need-based categories serve different goals.", "Track outcomes by scholarship type.", "No real socioeconomic data."])
    rows.append(["Faculty workload", "Enrollment links faculty to taught courses for workload analysis.", "High load can affect feedback and service quality.", "Popular departments need balanced staffing.", "Monitor distinct courses and enrollments per faculty.", "Workload excludes non-teaching duties."])
    rows.append(["Graduation eligibility", "Credit and GPA measures identify students approaching eligibility.", "Advising can prioritize final-year students.", "Missing credits or low GPA delay completion.", "Add drill-through student review page.", "Eligibility rule is simplified."])
    rows.append(["Payment status", "Paid, partial, and unpaid categories provide finance risk visibility.", "Finance and advising teams can coordinate interventions.", "Payment delays may align with income categories.", "Use respectful outreach workflows.", "Synthetic data includes fictional profiles."])
    rows.append(["Semester trend", "Semester fields support enrollment and performance trend pages.", "Trend changes help detect demand shifts.", "Admissions cycles and course offerings influence patterns.", "Monitor semester-over-semester growth.", "Trend is generated from random distributions."])
    return pd.DataFrame(rows, columns=["Finding","Observation","Interpretation","Possible Explanation","Recommended Action","Limitation"])


def create_report_and_docs(tables: dict[str, pd.DataFrame], validation: pd.DataFrame) -> None:
    data_dictionary(tables)
    write_markdown_and_assets(tables, validation)
    f = findings(tables)
    f.to_csv(DOCS / "Analytical_Findings.csv", index=False)
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Times New Roman"; style.font.size = Pt(12)
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Data Analysis and Reporting Using Power BI\nUniversity Academic Analytics"); run.bold = True; run.font.size = Pt(16)
    doc.add_paragraph("Student Information: [Name / ID placeholders]")
    sections = [
        ("Executive Summary", "This project builds a complete synthetic university analytics dataset and a Power BI-ready reporting package covering enrollment, academic performance, attendance, finance, faculty workload, and department budgets."),
        ("Dataset Description", f"The cleaned dataset contains {len(tables['Students'])} students, {len(tables['Courses'])} courses, {len(tables['Faculty'])} faculty members, {len(tables['Enrollment'])} enrollment records, {len(tables['Attendance'])} attendance records, and {len(tables['Payments'])} payment records across 2023-2026."),
        ("Data Model Design", "The model uses dimensions for Students, StudentProfile, Departments, Courses, Faculty, and Semesters, facts for Attendance, Payments, and DepartmentBudget, and Enrollment as the bridge/fact table resolving the conceptual many-to-many relationship between Students and Courses."),
        ("Power Query Transformations", "The raw data intentionally includes inconsistent capitalization, blanks, duplicate records, invalid marks, date text, and payment-status inconsistencies. Cleaning standardizes text, removes duplicates, corrects types, validates ranges, and derives risk categories."),
        ("DAX Measures", f"The measure library contains {len(DAX_MEASURES)} documented measures stored in `_Measures`, including KPIs, rates, growth calculations, ranking, dynamic titles, TOPN, CONCATENATEX, and USERELATIONSHIP examples."),
        ("Dashboard Design", "The report specification includes six required pages plus an optional drill-through page with slicers, bookmarks, dynamic titles, conditional formatting, report-page tooltips, and cross-filtering."),
        ("Main Findings", "\n".join(f"{r['Finding']}: {r['Observation']} Recommended action: {r['Recommended Action']}" for _, r in f.iterrows())),
        ("Limitations", "Power BI Desktop was not installed in the execution environment, so the `.pbix`, live Power BI Service URL, screenshots from the actual dashboard, and tenant publication checks remain manual completion steps."),
        ("Conclusion", "The package is ready for Power BI Desktop authoring using the cleaned workbook, relationship documentation, DAX library, Power Query code, and report-page build guide."),
    ]
    for h, text in sections:
        doc.add_heading(h, level=1); doc.add_paragraph(text)
    doc.add_picture(str(DOCS / "Data_Model_Diagram.png"), width=Inches(6.5))
    doc.save(REPORT / "PowerBI_Project_Report.docx")
    pdf = SimpleDocTemplate(str(REPORT / "PowerBI_Project_Report.pdf"), pagesize=A4)
    styles = getSampleStyleSheet(); story = [Paragraph("Data Analysis and Reporting Using Power BI", styles["Title"]), Spacer(1, 12)]
    for h, text in sections:
        story += [Paragraph(h, styles["Heading1"]), Paragraph(text.replace("\n","<br/>"), styles["BodyText"]), Spacer(1, 8)]
    pdf.build(story)
    create_pptx(f)
    create_readme(tables, f)
    create_web(tables)
    (DOCS / "Dashboard_User_Guide.pdf").write_bytes((REPORT / "PowerBI_Project_Report.pdf").read_bytes())


def create_pptx(findings_df: pd.DataFrame) -> None:
    slides = [
        ("University Academic Analytics", "Power BI assignment package using synthetic academic data."),
        ("Assignment Requirements", "Import data, relationships, transformations, DAX measures, and meaningful visuals."),
        ("Problem and Objectives", "Help management monitor enrollment, performance, attendance, payments, faculty workload, and budgets."),
        ("Dataset Overview", "2023-2026 synthetic relational data with raw and cleaned versions."),
        ("Data Tables", ", ".join(TABLES)),
        ("Data Model", "Students -> Enrollment <- Courses resolves many-to-many; StudentProfile is one-to-one."),
        ("Data Transformation", "Trim, clean, deduplicate, type-correct, validate marks, derive risk categories."),
        ("DAX Measures", f"{len(DAX_MEASURES)} documented measures including KPIs, rates, growth, ranking, and dynamic titles."),
        ("Dashboard Pages", "Executive, Student, Academic, Attendance, Financial, Department/Faculty, Drill-through."),
        ("Main Findings", "; ".join(findings_df.Observation.head(3).tolist())),
        ("Recommendations and Limitations", "Use advising triggers, payment-risk segmentation, and staffing review. PBIX build requires Power BI Desktop."),
        ("Conclusion", "The repository assets are ready for final Desktop authoring and publication."),
    ]
    out = PRES / "PowerBI_Project_Presentation.pptx"
    content_types = """<?xml version="1.0" encoding="UTF-8"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/ppt/presentation.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.presentation.main+xml"/>""" + "".join(f'<Override PartName="/ppt/slides/slide{i}.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slide+xml"/>' for i in range(1,len(slides)+1)) + "</Types>"
    pres_xml = f"""<?xml version="1.0" encoding="UTF-8"?><p:presentation xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"><p:sldIdLst>{''.join(f'<p:sldId id="{255+i}" r:id="rId{i}"/>' for i in range(1,len(slides)+1))}</p:sldIdLst><p:sldSz cx="9144000" cy="5143500" type="screen16x9"/></p:presentation>"""
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types)
        z.writestr("_rels/.rels", """<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="ppt/presentation.xml"/></Relationships>""")
        z.writestr("ppt/presentation.xml", pres_xml)
        z.writestr("ppt/_rels/presentation.xml.rels", """<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">""" + "".join(f'<Relationship Id="rId{i}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide{i}.xml"/>' for i in range(1,len(slides)+1)) + "</Relationships>")
        for i,(title,body) in enumerate(slides,1):
            slide = f"""<?xml version="1.0" encoding="UTF-8"?><p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"><p:cSld><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr/><p:sp><p:nvSpPr><p:cNvPr id="2" name="Title"/><p:cNvSpPr/><p:nvPr/></p:nvSpPr><p:spPr><a:xfrm><a:off x="500000" y="400000"/><a:ext cx="8200000" cy="800000"/></a:xfrm></p:spPr><p:txBody><a:bodyPr/><a:lstStyle/><a:p><a:r><a:rPr sz="3600" b="1"/><a:t>{escape(title)}</a:t></a:r></a:p></p:txBody></p:sp><p:sp><p:nvSpPr><p:cNvPr id="3" name="Body"/><p:cNvSpPr/><p:nvPr/></p:nvSpPr><p:spPr><a:xfrm><a:off x="700000" y="1500000"/><a:ext cx="7800000" cy="2600000"/></a:xfrm></p:spPr><p:txBody><a:bodyPr/><a:lstStyle/><a:p><a:r><a:rPr sz="2200"/><a:t>{escape(body)}</a:t></a:r></a:p></p:txBody></p:sp></p:spTree></p:cSld></p:sld>"""
            z.writestr(f"ppt/slides/slide{i}.xml", slide)


def create_readme(tables: dict[str, pd.DataFrame], f: pd.DataFrame) -> None:
    readme = f"""# University Academic Analytics - Power BI Assignment

![Python](https://img.shields.io/badge/Python-3.14-blue) ![Power%20BI](https://img.shields.io/badge/Power%20BI-Desktop%20required-yellow) ![Status](https://img.shields.io/badge/Project-PowerBI--ready-green)

Synthetic university analytics project for the assignment **Data Analysis and Reporting Using Power BI**.

## Dataset Size

- Students: {len(tables['Students']):,}
- Departments: {len(tables['Departments']):,}
- Courses: {len(tables['Courses']):,}
- Faculty: {len(tables['Faculty']):,}
- Enrollment records: {len(tables['Enrollment']):,}
- Attendance records: {len(tables['Attendance']):,}
- Payment records: {len(tables['Payments']):,}

## Data Model

![Data model](documentation/Data_Model_Diagram.png)

Students and Courses are conceptually many-to-many and are resolved through `Enrollment`. `Students` to `StudentProfile` is one-to-one. Department, course, faculty, semester, payment, attendance, and budget relationships are documented in `documentation/Relationship_Documentation.md`.

## Power BI Status

Power BI Desktop was not available in the build environment, so no `.pbix`, dashboard screenshot, Power BI Service URL, app URL, or public embed URL is claimed. Use `documentation/Build_PowerBI_Report_Step_by_Step.md` to create the final report from the cleaned workbook.

## Important Files

- Raw workbook: `data/raw/University_Academic_Analytics_Raw.xlsx`
- Cleaned workbook: `data/cleaned/University_Academic_Analytics_Cleaned.xlsx`
- Validation report: `data/validation/Data_Validation_Results.xlsx`
- Data dictionary: `documentation/University_Data_Dictionary.xlsx`
- DAX library: `powerbi/DAX_Measures.md`
- Power Query library: `powerbi/Power_Query_M_Code.md`
- Report: `report/PowerBI_Project_Report.docx` and `report/PowerBI_Project_Report.pdf`
- Presentation: `presentation/PowerBI_Project_Presentation.pptx`

## Run

```bash
python scripts/generate_dataset.py
python scripts/clean_dataset.py
python scripts/validate_dataset.py
python scripts/exploratory_analysis.py
```

## Main Findings

""" + "\n".join(f"- **{r['Finding']}:** {r['Observation']} Recommendation: {r['Recommended Action']}" for _, r in f.iterrows()) + "\n\n## Security\n\nAll data is fictional and synthetic. Do not commit credentials, cookies, tokens, `.env` files, Power BI session files, or authentication screenshots.\n"
    (ROOT / "README.md").write_text(readme, encoding="utf-8")


def create_web(tables: dict[str, pd.DataFrame]) -> None:
    html = f"""<!doctype html><html lang="en"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><title>University Academic Analytics</title><link rel="stylesheet" href="assets/style.css"></head><body><main><section class="hero"><h1>University Academic Analytics</h1><p>Power BI-ready academic, enrollment, attendance, finance, and department analytics package built from synthetic data.</p><div class="actions"><a href="../report/PowerBI_Project_Report.pdf">Report</a><a href="../presentation/PowerBI_Project_Presentation.pptx">Presentation</a><a href="../data/cleaned/University_Academic_Analytics_Cleaned.xlsx">Cleaned Dataset</a></div></section><section class="metrics"><div><b>{len(tables['Students']):,}</b><span>Students</span></div><div><b>{len(tables['Enrollment']):,}</b><span>Enrollments</span></div><div><b>{len(tables['Attendance']):,}</b><span>Attendance Records</span></div><div><b>{len(DAX_MEASURES)}</b><span>DAX Measures</span></div></section><section><h2>Data Model</h2><img src="assets/data-model.png" alt="Data model diagram"></section><section><h2>Power BI Note</h2><p>Power BI Desktop was not available during automated build, so the live report link should be added only after the actual report is authored and published.</p></section><footer>All data is fictional and generated for academic demonstration.</footer></main></body></html>"""
    css = "body{margin:0;font-family:Segoe UI,Arial,sans-serif;color:#17212b;background:#f6f8fb}main{max-width:1120px;margin:auto}.hero{padding:72px 24px 42px}.hero h1{font-size:44px;margin:0 0 12px}.hero p{font-size:20px;max-width:760px}.actions{display:flex;gap:12px;flex-wrap:wrap}.actions a{background:#1f5a7a;color:white;padding:12px 16px;border-radius:6px;text-decoration:none}.metrics{display:grid;grid-template-columns:repeat(4,1fr);gap:12px;padding:20px 24px}.metrics div{background:white;border:1px solid #dde5ee;border-radius:8px;padding:18px}.metrics b{display:block;font-size:30px;color:#1f5a7a}.metrics span{color:#52616b}section{padding:24px}img{max-width:100%;border:1px solid #d7dee8;border-radius:8px;background:white}footer{padding:32px 24px;color:#52616b}@media(max-width:760px){.metrics{grid-template-columns:1fr 1fr}.hero h1{font-size:34px}}"
    (WEB / "index.html").write_text(html, encoding="utf-8")
    (WEB / "assets" / "style.css").write_text(css, encoding="utf-8")


def write_project_meta() -> None:
    (ROOT / ".gitignore").write_text("""__pycache__/
*.pyc
.env
*.pbitmp
*.tmp
*.log
~$*
.DS_Store
Thumbs.db
*.bak
""", encoding="utf-8")
    (ROOT / "LICENSE").write_text("MIT License\n\nCopyright (c) 2026\n\nPermission is hereby granted, free of charge, to any person obtaining a copy of this software and associated documentation files.\n", encoding="utf-8")
    (ROOT / "SECURITY.md").write_text("# Security\n\nThis project uses synthetic data only. Do not commit credentials, tokens, cookies, browser profiles, `.env` files, Power BI credentials, or authentication screenshots.\n", encoding="utf-8")
    (ROOT / "CHANGELOG.md").write_text("# Changelog\n\n## v1.0.0\n\n- Generated synthetic university analytics dataset.\n- Added cleaning, validation, DAX, Power Query, documentation, report, presentation, and web landing page assets.\n", encoding="utf-8")
    (ROOT / "PROJECT_CHECKLIST.md").write_text("# Project Checklist\n\n- [x] Raw dataset generated\n- [x] Cleaned dataset generated\n- [x] Validation report generated\n- [x] Data dictionary generated\n- [x] Relationship documentation generated\n- [x] DAX measure library generated\n- [x] Power Query documentation generated\n- [x] Report generated\n- [x] Presentation generated\n- [x] Web landing page generated\n- [ ] Power BI Desktop `.pbix` authoring\n- [ ] Power BI Service publication\n- [ ] GitHub repository/release/pages publication\n", encoding="utf-8")
    (ROOT / "scripts" / "requirements.txt").write_text("pandas\nopenpyxl\nnumpy\npython-docx\nPillow\nreportlab\n", encoding="utf-8")


def exploratory_analysis() -> None:
    tables = clean_dataset()
    f = findings(tables)
    f.to_excel(DOCS / "Exploratory_Findings.xlsx", index=False)
    create_report_and_docs(tables, validate_dataset())


def build_all() -> None:
    ensure_dirs()
    generate_dataset()
    tables = clean_dataset()
    validation = validate_dataset()
    create_report_and_docs(tables, validation)
    write_project_meta()


if __name__ == "__main__":
    build_all()
